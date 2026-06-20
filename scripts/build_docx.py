import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / 'docs'
DATA_DICTIONARY = DOCS / 'data_dictionary.md'
DEPENDENCY_MATRIX = DOCS / 'dependency_matrix.md'
LINEAGE = DOCS / 'lineage.mmd'
OUTPUT = DOCS / 'documentacion.docx'

DARK_BLUE = RGBColor(31, 78, 121)
LIGHT_BLUE = RGBColor(46, 117, 182)
TEXT = RGBColor(37, 37, 37)
HEADING_RE = re.compile(r'^(#{1,6})\s+(.*)$')
LIST_RE = re.compile(r'^([\-*•])\s+(.*)$')
NUMBER_RE = re.compile(r'^\d+\.\s+(.*)$')


def clean_inline(text):
    text = text.replace('\u2022', '-')
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    text = text.replace('`', '')
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'__(.*?)__', r'\1', text)
    text = re.sub(r'(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)', r'\1', text)
    text = text.replace('  ', ' ')
    return text.strip()


def add_heading(doc, text, level=1, color=None):
    style = f'Heading {level}'
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(clean_inline(text))
    if color is not None:
        run.font.color.rgb = color
    return p


def add_paragraph(doc, text, style=None, color=None, bold=False, size=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = p.add_run(clean_inline(text))
    if color is not None:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    if size is not None:
        run.font.size = size
    return p


def parse_markdown_table(lines, start_idx):
    rows = []
    idx = start_idx
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        raw = lines[idx]
        if '---' in raw or raw.strip().startswith('|---'):
            idx += 1
            continue
        cells = [c.strip() for c in raw.strip().strip('|').split('|')]
        rows.append(cells)
        idx += 1
    return rows, idx


def render_table(doc, rows):
    if not rows:
        return

    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = 'Table Grid'
    table.autofit = True

    header = table.rows[0].cells
    for i, value in enumerate(rows[0]):
        header[i].text = clean_inline(value)
        for paragraph in header[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = clean_inline(value)
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def render_markdown_file(doc, path):
    lines = path.read_text(encoding='utf-8').splitlines()
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        heading_match = HEADING_RE.match(stripped)

        if heading_match:
            level = len(heading_match.group(1))
            text = clean_inline(heading_match.group(2))
            if level == 1:
                add_heading(doc, text, level=1, color=DARK_BLUE)
            elif level == 2:
                add_heading(doc, text, level=2, color=LIGHT_BLUE)
            else:
                add_heading(doc, text, level=min(level, 3), color=DARK_BLUE)
        elif stripped.startswith('|'):
            table_rows, i = parse_markdown_table(lines, i)
            if table_rows:
                render_table(doc, table_rows)
            continue
        elif stripped.startswith('> '):
            p = doc.add_paragraph(style='Intense Quote')
            p.add_run(clean_inline(stripped[2:]))
        elif LIST_RE.match(stripped):
            item = LIST_RE.match(stripped).group(2)
            add_paragraph(doc, item, style='List Bullet')
        elif NUMBER_RE.match(stripped):
            item = NUMBER_RE.match(stripped).group(1)
            add_paragraph(doc, item, style='List Number')
        elif stripped == '---':
            doc.add_paragraph()
        elif stripped:
            add_paragraph(doc, stripped)
        i += 1


def render_lineage(doc, path):
    add_heading(doc, 'Diagrama de lineage', level=1, color=DARK_BLUE)
    add_paragraph(doc, 'Diagrama Mermaid (versión textual)', color=TEXT)
    for raw in path.read_text(encoding='utf-8').splitlines():
        if raw.strip() and not raw.strip().startswith('%%'):
            p = doc.add_paragraph(style='Intense Quote')
            p.add_run(raw)


def build_docx():
    document = Document()
    for section in document.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run('Proyecto ADF Docs Lab')
    run_title.bold = True
    run_title.font.size = Pt(28)
    run_title.font.color.rgb = DARK_BLUE

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run('Documentación técnica consolidada')
    run_sub.bold = True
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = LIGHT_BLUE

    flow = document.add_paragraph()
    flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_flow = flow.add_run('Origen OLTP → Staging → Data Warehouse')
    run_flow.font.size = Pt(12)
    run_flow.font.color.rgb = TEXT

    document.add_page_break()

    add_heading(document, 'Documentación técnica consolidada', level=1, color=DARK_BLUE)
    add_paragraph(document, 'Este documento consolida la documentación derivada de los scripts SQL y del pipeline ADF.', color=TEXT)
    document.add_paragraph()

    render_markdown_file(document, DATA_DICTIONARY)
    document.add_page_break()
    render_markdown_file(document, DEPENDENCY_MATRIX)
    document.add_page_break()
    render_lineage(document, LINEAGE)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)

    # Validación real: el archivo debe abrirse correctamente con python-docx.
    from docx import Document as DocxDocument

    validated = DocxDocument(OUTPUT)
    print(f'Archivo generado y validado: {OUTPUT}')
    print(f'Parrafos={len(validated.paragraphs)} Tablas={len(validated.tables)}')


if __name__ == '__main__':
    build_docx()
